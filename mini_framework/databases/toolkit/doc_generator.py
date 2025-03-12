from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Pt
import os

from mini_framework.utils.log import logger


class StyleManager:
    def __init__(self, styles=None):
        logger.debug("Initializing StyleManager with styles: %s", styles)
        self.styles = styles or {
            "title": {
                "font_size": 16,
                "bold": True,
                "font_name": "黑体",
                "alignment": WD_PARAGRAPH_ALIGNMENT.CENTER,
            },
            "author": {
                "font_size": 12,
                "font_name": "宋体",
                "alignment": WD_PARAGRAPH_ALIGNMENT.CENTER,
            },
            "table_title": {"font_size": 14, "bold": True, "font_name": "宋体"},
            "header": {"font_size": 12, "bold": True, "font_name": "宋体"},
            "cell": {"font_size": 12, "font_name": "宋体"},
        }

    def apply_style(self, run, style_name):
        logger.debug("Applying style '%s' to run: %s", style_name, run.text)
        style = self.styles.get(style_name, {})
        run.font.size = Pt(style.get("font_size", 12))
        run.bold = style.get("bold", False)
        run.font.name = style.get("font_name", "宋体")
        if "font_name" in style:
            run._element.rPr.rFonts.set(qn("w:eastAsia"), style["font_name"])


class DocumentStructure:
    def __init__(
        self,
        title="数据库模型文档",
        file_name: str = None,
        author=None,
        table_config=None,
    ):
        logger.debug(
            "Initializing DocumentStructure with title: '%s', author: '%s'",
            title,
            author,
        )
        self.title = title
        self.file_name = file_name or f"{title}.docx"
        self.author = author
        self.table_config = table_config or {
            "title": "表结构信息",
            "headers": [
                "列名",
                "类型",
                "长度",
                "是否主键",
                "是否允许为空",
                "默认值",
                "列注释",
            ],
        }


class DocGenerator:
    def __init__(
        self,
        metadata_cls,
        output: str,
        structure: DocumentStructure,
        style_manager: StyleManager = None,
    ):
        """
        Initializes the document generator.
        :param metadata_cls: The metadata class.
        :param output: The output directory of the generated document.
        :param structure: An instance of DocumentStructure.
        :param style_manager: An instance of StyleManager.
        """
        logger.debug("Initializing DocGenerator with output: '%s'", output)
        self.__metadata_cls = metadata_cls
        self.__output = output
        self.__structure = structure
        self.__style_manager = style_manager or StyleManager()

    def create_word_document(self):
        """
        Creates a Word document.
        :return: The path of the generated document.
        """
        logger.info("Creating Word document")
        # Create a new Document
        doc = Document()

        # Set page orientation to landscape
        section = doc.sections[0]
        section.orientation = WD_ORIENT.LANDSCAPE
        new_width, new_height = section.page_height, section.page_width
        section.page_width = new_width
        section.page_height = new_height

        # Add title
        if self.__structure.title:
            logger.debug("Adding title: '%s'", self.__structure.title)
            title_para = doc.add_paragraph()
            title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run = title_para.add_run(self.__structure.title)
            self.__style_manager.apply_style(run, "title")

        # Add author
        if self.__structure.author:
            logger.debug("Adding author: '%s'", self.__structure.author)
            author_para = doc.add_paragraph()
            author_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run = author_para.add_run(f"作者: {self.__structure.author}")
            self.__style_manager.apply_style(run, "author")

        # Add table for metadata
        table_config = self.__structure.table_config
        if table_config:
            for table in self.__metadata_cls.tables.values():
                # Add spacing before each table
                doc.add_paragraph("\n")

                # Add table name and comment
                logger.debug(
                    "Adding table name: '%s', comment: '%s'", table.name, table.comment
                )
                table_name_para = doc.add_paragraph()
                table_name_para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                run = table_name_para.add_run(
                    f"表名: {table.name} ({table.comment or '无注释'})"
                )
                self.__style_manager.apply_style(run, "table_title")

                # Extract metadata and create a table
                headers = table_config.get("headers", [])
                logger.debug("Table headers: %s", headers)
                table_doc = doc.add_table(rows=1, cols=len(headers))
                table_doc.style = "Table Grid"

                # Set header row
                hdr_cells = table_doc.rows[0].cells
                for i, header in enumerate(headers):
                    logger.debug("Adding header cell: '%s'", header)
                    hdr_cells[i].text = header
                    for paragraph in hdr_cells[i].paragraphs:
                        run = paragraph.runs[0]
                        self.__style_manager.apply_style(run, "header")

                # Add metadata rows
                for column in table.columns:
                    logger.debug("Adding metadata for column: %s", column.name)
                    row_cells = table_doc.add_row().cells
                    row_cells[0].text = column.name
                    row_cells[1].text = str(column.type)

                    # Handle length or precision for floating-point types
                    if hasattr(column.type, "precision") and hasattr(
                        column.type, "scale"
                    ):
                        row_cells[2].text = (
                            f"精度: {column.type.precision}, 小数位: {column.type.scale}"
                        )
                    else:
                        row_cells[2].text = (
                            str(getattr(column.type, "length", "")) or ""
                        )

                    row_cells[3].text = "是" if column.primary_key else "否"
                    row_cells[4].text = "是" if column.nullable else "否"
                    row_cells[5].text = str(
                        column.default.arg if column.default else ""
                    )
                    row_cells[6].text = column.comment or ""

                    for cell in row_cells:
                        for paragraph in cell.paragraphs:
                            run = paragraph.runs[0]
                            self.__style_manager.apply_style(run, "cell")

        # Save document
        logger.debug("Saving document to output directory: '%s'", self.__output)
        os.makedirs(self.__output, exist_ok=True)
        output_path = os.path.join(self.__output, self.__structure.file_name)
        doc.save(output_path)

        logger.info("Document created successfully: '%s'", output_path)
        return output_path
